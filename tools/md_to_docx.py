"""Convert markdown to docx (basic formatting).

- Preserves headings (#), unordered/ordered lists, code blocks (```), and paragraphs.
- Leaves inline math/LaTeX and tables as raw text.

Usage:
    python tools/md_to_docx.py --input readme/TOTALREADME_katex_fixed.md --output readme/TOTALREADME_katex_fixed.docx
"""
import argparse
import os
import sys
import re
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn


def set_run_font(run, name=None, size_pt=None):
    font = run.font
    if name:
        font.name = name
        try:
            run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
        except Exception:
            pass
    if size_pt:
        font.size = Pt(size_pt)


def add_code_paragraph(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, name='Courier New', size_pt=9)


def convert(md_path, docx_path):
    doc = Document()
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    code_mode = False
    code_lines = []

    for raw in lines:
        line = raw.rstrip('\n')
        stripped = line.strip()

        # Toggle code fence
        if stripped.startswith('```'):
            code_mode = not code_mode
            if not code_mode:
                # flush code block
                if code_lines:
                    for cl in code_lines:
                        add_code_paragraph(doc, cl)
                    code_lines = []
            continue

        if code_mode:
            code_lines.append(line)
            continue

        # Blank line -> paragraph break
        if stripped == '':
            doc.add_paragraph()
            continue

        # Headings
        if stripped.startswith('#'):
            m = re.match(r'^(#+)\s*(.*)$', stripped)
            if m:
                level = min(len(m.group(1)), 4)
                text = m.group(2).strip()
                try:
                    # docx supports levels 0..9, where 0 is Title
                    doc.add_heading(text, level=level)
                except Exception:
                    p = doc.add_paragraph()
                    r = p.add_run(text)
                    r.bold = True
                continue

        # Unordered list
        m = re.match(r'^\s*[-\*]\s+(.*)$', line)
        if m:
            try:
                doc.add_paragraph(m.group(1).strip(), style='List Bullet')
            except Exception:
                doc.add_paragraph('\u2022 ' + m.group(1).strip())
            continue

        # Ordered list
        m = re.match(r'^\s*(\d+)\.\s+(.*)$', line)
        if m:
            try:
                doc.add_paragraph(m.group(2).strip(), style='List Number')
            except Exception:
                doc.add_paragraph(m.group(2).strip())
            continue

        # Default paragraph
        doc.add_paragraph(line)

    # If file ended while still in code block, flush
    if code_mode and code_lines:
        for cl in code_lines:
            add_code_paragraph(doc, cl)

    doc.save(docx_path)
    print(f"Wrote {docx_path}")


if __name__ == '__main__':
    parser = argparse.ArgumentParser()
    parser.add_argument('--input', '-i', default='readme/TOTALREADME_katex_fixed.md')
    parser.add_argument('--output', '-o', default='readme/TOTALREADME_katex_fixed.docx')
    args = parser.parse_args()

    md = args.input
    out = args.output
    if not os.path.exists(md):
        print(f"Input file not found: {md}", file=sys.stderr)
        sys.exit(2)
    try:
        convert(md, out)
    except Exception as e:
        print("Conversion failed:", e, file=sys.stderr)
        sys.exit(1)
    sys.exit(0)
