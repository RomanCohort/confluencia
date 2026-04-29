"""Report Generator for ConfluenciaStudio.

Generates publication-ready reports in Markdown/LaTeX format
from experiment results and analysis.
"""

from __future__ import annotations

import json
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional, Any
from dataclasses import dataclass, field


@dataclass
class ReportSection:
    """A section in the report."""
    title: str
    content: str
    subsections: List['ReportSection'] = field(default_factory=list)
    figures: List[str] = field(default_factory=list)  # Figure paths
    tables: List[Dict] = field(default_factory=list)


class BioinformaticsReportGenerator:
    """Generates Bioinformatics journal format reports.

    Features:
    - Structured sections (Abstract, Methods, Results, etc.)
    - Table generation from dataframes
    - Figure embedding
    - LaTeX and Markdown output
    - Bibliography formatting

    Usage:
        generator = BioinformaticsReportGenerator()
        report = generator.generate_draft(experiment, sections=["methods", "results"])
        generator.export_to_latex(report, "output/report.tex")
    """

    def __init__(self, template: str = "bioinformatics"):
        self.template = template
        self.sections_order = [
            "abstract",
            "introduction",
            "methods",
            "results",
            "discussion",
            "conclusion",
            "references",
        ]

    def generate_draft(
        self,
        experiment_data: Dict[str, Any],
        sections: Optional[List[str]] = None,
        title: Optional[str] = None,
        authors: Optional[List[str]] = None,
    ) -> str:
        """Generate a report draft from experiment data.

        Args:
            experiment_data: Dict with experiment info, metrics, parameters
            sections: List of sections to include (default: all)
            title: Report title
            authors: List of author names

        Returns:
            Markdown formatted report
        """
        sections = sections or self.sections_order
        title = title or experiment_data.get("name", "Computational Analysis Report")

        lines = []

        # Title
        lines.append(f"# {title}\n")

        # Authors
        if authors:
            lines.append(f"**Authors:** {', '.join(authors)}\n")
            lines.append(f"**Date:** {datetime.now().strftime('%Y-%m-%d')}\n")
            lines.append("---\n")

        # Generate each section
        for section in sections:
            section_method = getattr(self, f"_generate_{section}", None)
            if section_method:
                content = section_method(experiment_data)
                lines.append(content)
                lines.append("\n")

        return "\n".join(lines)

    def _generate_abstract(self, data: Dict) -> str:
        """Generate abstract section."""
        lines = ["## Abstract\n"]

        # Extract key metrics for summary
        metrics = data.get("metrics", {})
        module = data.get("module", "analysis")

        if module == "drug":
            r2 = metrics.get("r2", metrics.get("R2", "N/A"))
            rmse = metrics.get("rmse", metrics.get("RMSE", "N/A"))
            lines.append(
                f"We developed a machine learning model for drug efficacy prediction. "
                f"The model achieved R² = {r2} and RMSE = {rmse} on the test set. "
                f"This tool enables rapid screening of drug candidates for circRNA-based therapeutics.\n"
            )
        elif module == "epitope":
            auc = metrics.get("auc", metrics.get("AUC", "N/A"))
            lines.append(
                f"We developed an epitope-MHC binding prediction model. "
                f"The model achieved AUC = {auc} on the validation set. "
                f"This tool supports vaccine design by identifying immunogenic epitopes.\n"
            )
        else:
            lines.append(
                "Computational analysis was performed using the Confluencia platform. "
                "Results demonstrate the utility of machine learning approaches in circRNA drug discovery.\n"
            )

        return "".join(lines)

    def _generate_introduction(self, data: Dict) -> str:
        """Generate introduction section."""
        lines = ["## Introduction\n"]

        module = data.get("module", "")

        if module == "drug":
            lines.append("""
Circular RNAs (circRNAs) have emerged as promising therapeutic candidates due to their stability
and tissue-specific expression patterns. Predicting drug efficacy for circRNA-based therapeutics
remains a significant challenge in drug development.

Machine learning approaches offer a powerful alternative to traditional pharmacokinetic modeling,
enabling rapid screening of large compound libraries. In this study, we develop and validate
a prediction model using Mixture of Experts (MOE) ensemble learning.

**Objectives:**
1. Develop a predictive model for circRNA drug efficacy
2. Validate the model using held-out test data
3. Compare performance across different model architectures

""")
        elif module == "epitope":
            lines.append("""
Epitope-MHC binding prediction is crucial for vaccine design and cancer immunotherapy.
Accurate computational prediction can dramatically reduce experimental screening costs.

Deep learning and ensemble methods have shown promising results for this task. We present
a framework combining sequence encoding, MHC pseudo-sequence features, and ensemble learning.

**Objectives:**
1. Train models on IEDB benchmark data
2. Incorporate MHC allele-specific features
3. Achieve competitive AUC performance

""")
        else:
            lines.append("""
This report presents computational analysis results from the Confluencia platform
for circRNA drug discovery. The analysis encompasses multiple computational approaches
including machine learning, pharmacokinetic simulation, and multi-omics integration.

""")

        return "".join(lines)

    def _generate_methods(self, data: Dict) -> str:
        """Generate methods section."""
        lines = ["## Methods\n"]

        params = data.get("parameters", {})
        module = data.get("module", "")

        # Dataset
        lines.append("### Dataset\n")
        data_path = params.get("data_path", params.get("data", "N/A"))
        lines.append(f"- Training data: `{data_path}`\n")

        if "test_size" in params:
            lines.append(f"- Test set proportion: {params['test_size']}\n")

        # Model architecture
        lines.append("\n### Model Architecture\n")
        model_type = params.get("model_type", params.get("model-type", "MOE"))
        lines.append(f"- Model type: **{model_type}**\n")

        if "n_estimators" in params:
            lines.append(f"- Number of estimators: {params['n_estimators']}\n")

        # Features
        lines.append("\n### Feature Engineering\n")
        if module == "drug":
            lines.append("""
- Morgan fingerprints (2048 bits, radius 2)
- Physicochemical descriptors (MW, LogP, TPSA, HBD, HBA)
- Dose-response features (Emax model parameters)
- Cross-feature interactions

""")
        elif module == "epitope":
            lines.append("""
- Amino acid composition (20 features)
- Dipeptide composition (400 features)
- MHC pseudo-sequence encoding (when allele specified)
- ESM-2 embeddings (optional)

""")

        # Training procedure
        lines.append("### Training Procedure\n")
        lines.append("""
1. Data preprocessing: SMILES validation, missing value handling
2. Feature computation: Descriptor calculation, fingerprint generation
3. Model training: Cross-validation for hyperparameter selection
4. Evaluation: Test set prediction and metric computation

""")

        return "".join(lines)

    def _generate_results(self, data: Dict) -> str:
        """Generate results section."""
        lines = ["## Results\n"]

        metrics = data.get("metrics", {})
        params = data.get("parameters", {})

        # Performance metrics
        lines.append("### Model Performance\n")

        if metrics:
            lines.append("| Metric | Value |\n")
            lines.append("|--------|-------|\n")

            for key, value in sorted(metrics.items()):
                if isinstance(value, float):
                    lines.append(f"| {key.upper()} | {value:.4f} |\n")
                else:
                    lines.append(f"| {key.upper()} | {value} |\n")

            lines.append("\n")

        # Feature importance
        if "feature_importance" in data:
            lines.append("### Feature Importance\n")
            fi = data["feature_importance"]
            lines.append("Top contributing features:\n")
            for feat, imp in list(fi.items())[:10]:
                lines.append(f"- {feat}: {imp:.4f}\n")
            lines.append("\n")

        # Figures reference
        artifacts = data.get("artifacts", [])
        figures = [a for a in artifacts if any(ext in a.lower() for ext in ['.png', '.svg', '.pdf'])]

        if figures:
            lines.append("### Figures\n")
            for i, fig_path in enumerate(figures[:5], 1):
                lines.append(f"![Figure {i}]({fig_path})\n")
            lines.append("\n")

        return "".join(lines)

    def _generate_discussion(self, data: Dict) -> str:
        """Generate discussion section."""
        lines = ["## Discussion\n"]

        metrics = data.get("metrics", {})
        r2 = metrics.get("r2", metrics.get("R2", None))

        if r2 is not None:
            if r2 > 0.7:
                lines.append(f"""
The model achieved strong predictive performance with R² = {r2:.4f}, indicating
good generalization to unseen molecules. The Mixture of Experts architecture
effectively combines predictions from diverse base learners.

Key factors contributing to model performance:
- Comprehensive feature engineering including dose-response parameters
- Cross-feature interactions capturing non-linear relationships
- Ensemble averaging reducing prediction variance

""")
            else:
                lines.append(f"""
The model achieved moderate predictive performance with R² = {r2:.4f}.
This suggests potential for improvement through:

1. Additional feature engineering (GNN embeddings, ChemBERTa)
2. Larger training datasets
3. Fine-tuning of model hyperparameters

""")

        lines.append("""
**Limitations:**
- Model trained on limited molecular diversity
- External validation on independent datasets recommended
- Predictions should be confirmed experimentally

""")

        return "".join(lines)

    def _generate_conclusion(self, data: Dict) -> str:
        """Generate conclusion section."""
        lines = ["## Conclusion\n"]

        module = data.get("module", "analysis")

        lines.append(f"""
We have developed a computational pipeline for {module} prediction as part of the
Confluencia circRNA drug discovery platform. The results demonstrate the utility
of machine learning approaches in accelerating therapeutic development.

**Key Contributions:**
- Trained and validated prediction model
- Established benchmark performance metrics
- Provided reproducible analysis pipeline

**Future Work:**
- Expand training data with diverse molecular scaffolds
- Integrate multi-task learning for related endpoints
- Develop uncertainty quantification for predictions

""")

        return "".join(lines)

    def _generate_references(self, data: Dict) -> str:
        """Generate references section."""
        lines = ["## References\n"]

        lines.append("""
1. Chen, L.L. (2016). The biogenesis and emerging roles of circular RNAs. *Nature Reviews Molecular Cell Biology*, 17(4), 205-211.

2. Mayr, A., et al. (2016). DeepTox: Toxicity prediction using deep learning. *Frontiers in Environmental Science*, 3, 80.

3. O'Donnell, T.J., et al. (2018). MHCflurry: open-source class I MHC binding affinity prediction. *Cell Systems*, 7(1), 129-132.

4. Pedregosa, F., et al. (2011). Scikit-learn: Machine learning in Python. *JMLR*, 12, 2825-2830.

5. RDKit: Open-source cheminformatics. http://www.rdkit.org

""")

        return "".join(lines)

    def add_table(self, data: Dict[str, List], caption: str) -> str:
        """Generate a formatted table.

        Args:
            data: Dict with column names as keys and value lists
            caption: Table caption

        Returns:
            Markdown table string
        """
        if not data:
            return ""

        columns = list(data.keys())
        rows = len(data[columns[0]])

        lines = [f"**{caption}**\n"]
        lines.append("| " + " | ".join(columns) + " |")
        lines.append("| " + " | ".join(["---"] * len(columns)) + " |")

        for i in range(rows):
            values = [str(data[col][i]) if i < len(data[col]) else "" for col in columns]
            lines.append("| " + " | ".join(values) + " |")

        return "\n".join(lines) + "\n"

    def export_to_latex(self, markdown_content: str, output_path: str) -> str:
        """Convert markdown report to LaTeX format.

        Args:
            markdown_content: Markdown report content
            output_path: Output .tex file path

        Returns:
            Path to generated LaTeX file
        """
        latex_lines = [
            "\\documentclass{article}",
            "\\usepackage[utf8]{inputenc}",
            "\\usepackage{graphicx}",
            "\\usepackage{booktabs}",
            "\\usepackage{hyperref}",
            "",
            "\\title{Computational Analysis Report}",
            "\\author{Confluencia Studio}",
            "\\date{" + datetime.now().strftime('%Y-%m-%d') + "}",
            "",
            "\\begin{document}",
            "\\maketitle",
            "",
        ]

        # Convert markdown to LaTeX (simplified)
        for line in markdown_content.split('\n'):
            if line.startswith('# '):
                latex_lines.append(f"\\section{{{line[2:]}}}")
            elif line.startswith('## '):
                latex_lines.append(f"\\subsection{{{line[3:]}}}")
            elif line.startswith('### '):
                latex_lines.append(f"\\subsubsection{{{line[4:]}}}")
            elif line.startswith('**') and line.endswith('**'):
                latex_lines.append(f"\\textbf{{{line[2:-2]}}}")
            elif line.startswith('!['):
                # Extract image path
                import re
                match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line)
                if match:
                    latex_lines.append(f"\\includegraphics[width=0.8\\textwidth]{{{match.group(2)}}}")
            elif line.startswith('| '):
                # Table row
                cells = [c.strip() for c in line.split('|')[1:-1]]
                if all(c == '---' for c in cells):
                    latex_lines.append("\\hline")
                else:
                    latex_lines.append(" & ".join(cells) + " \\\\")
            elif line.startswith('- '):
                latex_lines.append(f"\\item {line[2:]}")
            elif line.strip():
                latex_lines.append(line)

        latex_lines.extend([
            "",
            "\\end{document}",
        ])

        latex_content = "\n".join(latex_lines)

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(latex_content)

        return output_path

    def export_to_docx(self, markdown_content: str, output_path: str) -> str:
        """Export report to DOCX format.

        Requires python-docx package.
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches

            doc = Document()

            # Parse and add content
            for line in markdown_content.split('\n'):
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('**') and line.endswith('**'):
                    p = doc.add_paragraph()
                    run = p.add_run(line[2:-2])
                    run.bold = True
                elif line.startswith('- '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif line.strip():
                    doc.add_paragraph(line)

            doc.save(output_path)
            return output_path

        except ImportError:
            raise ImportError("python-docx is required for DOCX export. Install with: pip install python-docx")
