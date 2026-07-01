#!/usr/bin/env python3
"""
generate_caspcirc_word.py - 生成Circ-CASP 2026参赛名单Word文档

安装依赖: pip install python-docx
运行: python generate_caspcirc_word.py
"""

def generate_word_document():
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # 创建文档
        doc = Document()

        # 标题
        title = doc.add_heading('Circ-CASP 2026 参赛名单', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 简介
        doc.add_paragraph('经过激烈（并没有）的报名环节，我们收到了来自全球（主要是我自己）的9份参赛申请。')
        doc.add_paragraph('经过严格（确实）的技术审核，9支队伍全部入围，现进行初步公布。')
        doc.add_paragraph('')

        # ===== 正式赛道 =====
        doc.add_heading('正式赛道', level=1)

        teams_main = [
            ('Team 1', '吉林大学计算机科学与技术学院', 'EGNN + 物理精修', 'Scheme 1', '已实现'),
            ('Team 2', '吉林大学计算机科学与技术学院', '原子力场求解', 'Scheme 2', '已实现'),
            ('Team 3', '吉林大学计算机科学与技术学院', '双引擎迭代蒸馏', 'Scheme 3', '已实现'),
            ('Team 4', '吉林大学计算机科学与技术学院', '坐标扩散 + EGNN', 'Scheme 4', '已实现'),
            ('Team 5', '吉林大学计算机科学与技术学院', 'Transformer 物理bias', 'Scheme 5', '已弃用'),
            ('Team 6', '吉林大学计算机科学与技术学院', '隐空间扩散', 'Scheme 6', '已实现'),
            ('Team 7', '吉林大学计算机科学与技术学院', '局部注意力+环式Mamba', 'Scheme 7', '⭐推荐'),
            ('Team 8', '吉林大学计算机科学与技术学院', '稀疏配对引导混合', 'Scheme 8', '已实现'),
            ('Team 9', '吉林大学计算机科学与技术学院', '线性RNA环化', 'Scheme 0', '官方基线'),
        ]

        for team_id, institution, method, scheme, status in teams_main:
            p = doc.add_paragraph()
            p.add_run(f'{team_id}').bold = True
            p.add_run(f'（{institution}）：{method} ')
            run = p.add_run(f'[{scheme} - {status}]')
            run.italic = True

        doc.add_paragraph('')

        # ===== 神仙打架赛道 =====
        doc.add_heading('神仙打架赛道', level=1)

        teams_expert = [
            ('Team 10', '浙江大学生命科学学院和定量生物学研究所', 'isRNAcirc'),
            ('Team 11', '山东大学数学与交叉科学研究中心，经环化改进', 'trRosettaRNA2'),
            ('Team 12', '维也纳大学理论化学系，经坐标映射改进', 'ViennaRNA-Circ'),
        ]

        for team_id, institution, method in teams_expert:
            p = doc.add_paragraph()
            p.add_run(f'{team_id}').bold = True
            p.add_run(f'（方法来自{institution}）：{method}')

        doc.add_paragraph('')

        # ===== 随机数赛道 =====
        doc.add_heading('随机数赛道', level=1)

        p = doc.add_paragraph()
        p.add_run('Team 13').bold = True
        p.add_run(': {114514, 67, 886}')

        doc.add_paragraph('')

        # ===== Scheme对应表 =====
        doc.add_heading('Scheme编号对应关系', level=1)

        table = doc.add_table(rows=10, cols=4)
        table.style = 'Light Grid Accent 1'

        # 表头
        headers = ['Team', '方法名称', 'Scheme', '状态']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True

        # 数据行
        scheme_data = [
            ('Team 1', 'EGNN + 物理精修', 'Scheme 1', '已实现'),
            ('Team 2', '原子力场求解', 'Scheme 2', '已实现'),
            ('Team 3', '双引擎迭代蒸馏', 'Scheme 3', '已实现'),
            ('Team 4', '坐标扩散 + EGNN', 'Scheme 4', '已实现'),
            ('Team 5', 'Transformer 物理bias', 'Scheme 5', '已弃用'),
            ('Team 6', '隐空间扩散', 'Scheme 6', '已实现'),
            ('Team 7', '局部注意力+环式Mamba', 'Scheme 7', '推荐'),
            ('Team 8', '稀疏配对引导混合', 'Scheme 8', '已实现'),
            ('Team 9', '线性RNA环化', 'Scheme 0', '官方基线'),
        ]

        for i, (team, method, scheme, status) in enumerate(scheme_data):
            table.rows[i+1].cells[0].text = team
            table.rows[i+1].cells[1].text = method
            table.rows[i+1].cells[2].text = scheme
            table.rows[i+1].cells[3].text = status

        doc.add_paragraph('')

        # ===== Team 9特殊地位 =====
        doc.add_heading('Team 9的特殊地位', level=1)

        p = doc.add_paragraph()
        p.add_run('Team 9 = Scheme 0 = CircFold Baseline = 线性RNA环化法').bold = True

        roles = [
            '官方基线方法 - 所有其他队伍的benchmark',
            '数据生成器 - 为其他方法提供训练数据（8万条）',
            'Teacher模型 - 为Team 3提供知识蒸馏来源',
            'Pipeline - 5-stage物理优化流程（ViennaRNA → trRosettaRNA2 → OpenMM → MD → Filter）'
        ]

        for role in roles:
            doc.add_paragraph(role, style='List Bullet')

        doc.add_paragraph('')

        # ===== 预测排名 =====
        doc.add_heading('预测获胜排名', level=1)

        rankings = [
            ('🥇 第1名', 'Team 7', 'Mamba长距离依赖 + 环式优化，最适合circRNA'),
            ('🥈 第2名', 'Team 3', 'Team 9作为Teacher，知识蒸馏优势'),
            ('🥉 第3名', 'Team 9', '官方基线，物理优化保证质量'),
            ('第4名', 'Team 8', '稀疏配对引导，BSJ准确率高'),
            ('第5名', 'Team 10', '成熟外部方法（isRNAcirc）'),
        ]

        for medal, team, reason in rankings:
            p = doc.add_paragraph()
            p.add_run(f'{medal} - {team}：').bold = True
            p.add_run(reason)

        doc.add_paragraph('')

        # ===== 页脚 =====
        doc.add_paragraph('—' * 30)
        p = doc.add_paragraph()
        p.add_run('Circ-CASP 2026 - Advancing circRNA 3D Structure Prediction').italic = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 保存文档
        output_path = 'CircCASP_2026_Participants.docx'
        doc.save(output_path)

        print(f'\n✓ Word文档生成成功！')
        print(f'  文件名: {output_path}')
        print(f'  路径: {output_path}')

        return output_path

    except ImportError:
        print('\n❌ 缺少依赖库 python-docx')
        print('\n请运行以下命令安装:')
        print('  pip install python-docx')
        print('\n安装后重新运行此脚本即可生成Word文档。')
        return None


if __name__ == '__main__':
    print('='*70)
    print('Circ-CASP 2026 参赛名单 Word文档生成器')
    print('='*70)
    generate_word_document()
