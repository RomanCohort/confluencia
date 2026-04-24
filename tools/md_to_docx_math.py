"""Convert Markdown to DOCX and render LaTeX formulas to images.

- Renders display (`$$...$$`) and inline (`$...$` or `\(...\)`) math using matplotlib mathtext.
- Embeds display formulas as centered images and inline formulas as inline images when possible.

Usage:
  & ".venv\Scripts\python.exe" tools\md_to_docx_math.py --input readme/TOTALREADME_katex_fixed.md --output readme/TOTALREADME_katex_fixed_math.docx
"""

import argparse
import os
import re
import shutil
import tempfile
from pathlib import Path

try:
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    from matplotlib import font_manager as fm
except Exception as e:
    print('matplotlib is required for formula rendering:', e)
    raise

try:
    from docx import Document
    from docx.shared import Inches, Pt
except Exception as e:
    print('python-docx is required:', e)
    raise

# Prefer common Windows Chinese fonts for proper glyph rendering in formulas
try:
    preferred_fonts = ['SimSun', 'SimHei', 'Microsoft YaHei', 'Microsoft YaHei UI', 'Arial Unicode MS']
    available_font_names = {f.name for f in fm.fontManager.ttflist}
    for pf in preferred_fonts:
        if pf in available_font_names:
            plt.rcParams['font.sans-serif'] = [pf]
            plt.rcParams['font.family'] = 'sans-serif'
            break
    plt.rcParams['axes.unicode_minus'] = False
    # Use STIX for math if available to improve symbol set
    plt.rcParams['mathtext.fontset'] = 'stix'
except Exception:
    pass


def render_formula_image(formula, out_path, fontsize=18, dpi=300):
    # Render formula (LaTeX math) to PNG using matplotlib mathtext
    fig = plt.figure()
    fig.patch.set_alpha(0.0)
    text = fig.text(0, 0, f"${formula}$", fontsize=fontsize)
    # Draw to compute bounding box
    fig.canvas.draw()
    try:
        renderer = fig.canvas.get_renderer()
        bbox = text.get_window_extent(renderer=renderer)
    except Exception:
        bbox = text.get_window_extent()
    width, height = bbox.width / fig.dpi, bbox.height / fig.dpi
    # Add small padding
    pad_w, pad_h = 0.2, 0.2
    fig.set_size_inches(max(width + pad_w, 0.5), max(height + pad_h, 0.5))
    text.set_position((0.05, 0.05))
    fig.canvas.draw()
    fig.savefig(out_path, dpi=dpi, transparent=True, bbox_inches='tight', pad_inches=0.02)
    plt.close(fig)


def split_code_segments(md_text):
    segments = []
    cur = []
    in_code = False
    for line in md_text.splitlines(keepends=True):
        if line.strip().startswith('```'):
            if not in_code:
                if cur:
                    segments.append(('text', ''.join(cur)))
                    cur = []
                in_code = True
                cur.append(line)
            else:
                cur.append(line)
                segments.append(('code', ''.join(cur)))
                cur = []
                in_code = False
        else:
            cur.append(line)
    if cur:
        segments.append(('code' if in_code else 'text', ''.join(cur)))
    return segments


def extract_formulas(text):
    """Replace formulas with placeholders and return mapping placeholder -> (formula, display)"""
    mapping = {}
    counter = 1

    # display $$...$$
    def repl_display(m):
        nonlocal counter
        content = m.group(1).strip()
        key = f"<<FORMULA_ID_{counter}>>"
        mapping[key] = (content, True)
        counter += 1
        return key

    text = re.sub(r'\$\$(.+?)\$\$', repl_display, text, flags=re.DOTALL)

    # \( ... \)
    def repl_paren(m):
        nonlocal counter
        content = m.group(1).strip()
        key = f"<<FORMULA_ID_{counter}>>"
        mapping[key] = (content, False)
        counter += 1
        return key

    text = re.sub(r'\\\((.+?)\\\)', repl_paren, text, flags=re.DOTALL)

    # single-dollar inline (avoid $$)
    def repl_single(m):
        nonlocal counter
        content = m.group(1).strip()
        key = f"<<FORMULA_ID_{counter}>>"
        mapping[key] = (content, False)
        counter += 1
        return key

    text = re.sub(r'(?<!\$)\$(?!\$)(.+?)(?<!\$)\$(?!\$)', repl_single, text, flags=re.DOTALL)

    return text, mapping


def md_to_docx_with_math(md_path, docx_path, images_dir):
    os.makedirs(images_dir, exist_ok=True)
    with open(md_path, 'r', encoding='utf-8') as f:
        md = f.read()

    segments = split_code_segments(md)

    # Process text segments to replace formulas with placeholders
    processed_segments = []
    global_mapping = {}
    for typ, content in segments:
        if typ == 'text':
            new_text, mapping = extract_formulas(content)
            processed_segments.append(('text', new_text))
            global_mapping.update(mapping)
        else:
            processed_segments.append((typ, content))

    doc = Document()

    # iterate segments and write to docx, substituting placeholders with images
    formula_cache = {}

    def ensure_formula_image(key, formula, display):
        if key in formula_cache:
            return formula_cache[key]
        # sanitize filename
        fname = f"formula_{key.strip('<>').replace('FORMULA_ID_','')}.png"
        out_path = os.path.join(images_dir, fname)
        try:
            render_formula_image(formula, out_path, fontsize=18 if not display else 20)
        except Exception as e:
            # fallback: write plain text file
            with open(out_path, 'wb') as wf:
                wf.write(b'')
            print('Failed to render formula', formula, e)
        formula_cache[key] = out_path
        return out_path

    for typ, content in processed_segments:
        if typ == 'code':
            # write code block lines
            lines = content.splitlines()
            # strip the fence lines if present
            if lines and lines[0].strip().startswith('```'):
                lines = lines[1:]
            if lines and lines[-1].strip().startswith('```'):
                lines = lines[:-1]
            for ln in lines:
                p = doc.add_paragraph()
                r = p.add_run(ln.rstrip('\n'))
                r.font.name = 'Courier New'
                r.font.size = Pt(9)
            continue

        # handle text lines
        for raw in content.splitlines():
            line = raw.rstrip('\n')
            stripped = line.strip()

            # headings
            if stripped.startswith('#'):
                m = re.match(r'^(#+)\s*(.*)$', stripped)
                if m:
                    level = min(len(m.group(1)), 4)
                    text = m.group(2).strip()
                    try:
                        doc.add_heading(text, level=level)
                    except Exception:
                        p = doc.add_paragraph()
                        r = p.add_run(text)
                        r.bold = True
                    continue

            # unordered list
            m = re.match(r'^\s*[-\*]\s+(.*)$', line)
            if m:
                try:
                    doc.add_paragraph(m.group(1).strip(), style='List Bullet')
                except Exception:
                    doc.add_paragraph('\u2022 ' + m.group(1).strip())
                continue

            # ordered list
            m = re.match(r'^\s*(\d+)\.\s+(.*)$', line)
            if m:
                try:
                    doc.add_paragraph(m.group(2).strip(), style='List Number')
                except Exception:
                    doc.add_paragraph(m.group(2).strip())
                continue

            # empty line
            if stripped == '':
                doc.add_paragraph()
                continue

            # Process inline placeholders in this line
            parts = re.split(r'(<<FORMULA_ID_\d+>>)', line)
            if len(parts) == 1:
                doc.add_paragraph(line)
                continue

            p = doc.add_paragraph()
            for part in parts:
                if not part:
                    continue
                if part.startswith('<<FORMULA_ID_'):
                    if part not in global_mapping:
                        # unknown placeholder -> insert literal
                        r = p.add_run(part)
                        continue
                    formula, display = global_mapping[part]
                    img_path = ensure_formula_image(part, formula, display)
                    if display:
                        # finish current paragraph and insert centered image paragraph
                        # remove empty trailing paragraph if any
                        # add image as its own paragraph
                        # if current paragraph has text, keep it and then add image paragraph
                        # add image paragraph
                        img_p = doc.add_paragraph()
                        try:
                            img_p.alignment = 1  # center
                        except Exception:
                            pass
                        try:
                            doc.add_picture(img_path, width=Inches(4.5))
                        except Exception:
                            # fallback: add run with placeholder text
                            img_p.add_run('[formula]')
                    else:
                        # inline: attempt to add image inline
                        run = p.add_run()
                        if hasattr(run, 'add_picture'):
                            try:
                                run.add_picture(img_path, width=Inches(1.2))
                            except Exception:
                                # fallback to separate picture paragraph
                                doc.add_picture(img_path, width=Inches(1.2))
                        else:
                            # fallback to separate paragraph picture
                            doc.add_picture(img_path, width=Inches(1.2))
                else:
                    p.add_run(part)

    doc.save(docx_path)
    print('Wrote', docx_path)


if __name__ == '__main__':
    parser = argparse.ArgumentParser()
    parser.add_argument('--input', '-i', default='readme/TOTALREADME_katex_fixed.md')
    parser.add_argument('--output', '-o', default='readme/TOTALREADME_katex_fixed_math.docx')
    args = parser.parse_args()

    md = args.input
    out = args.output
    if not os.path.exists(md):
        print('Input file not found:', md)
        raise SystemExit(2)

    images_dir = os.path.join('tmp', 'formula_images')
    try:
        md_to_docx_with_math(md, out, images_dir)
    except Exception as e:
        print('Conversion failed:', e)
        raise
